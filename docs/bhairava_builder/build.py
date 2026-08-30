from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT
import re
import copy
import os

TEMPLATE = "template.docx"
INPUT_DIR = "input"
OUTPUT_DIR = "output"
LOG_FILE = "logs.txt"

os.makedirs(OUTPUT_DIR, exist_ok=True)


def is_devanagari(text: str) -> bool:
    return any("\u0900" <= ch <= "\u097F" for ch in text)


def is_iast(text: str) -> bool:
    return bool(re.search(r"[āīūṛṝḷḹṅñṭḍṇśṣḥṃṁ]", text))


def is_number_line(text: str) -> bool:
    text = text.strip()
    return bool(re.fullmatch(r"[०-९]+", text))


def get_style(doc: Document, names):
    lowered = {n.lower() for n in names}
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.lower() in lowered:
            return p.style
    for style in doc.styles:
        if style.name and style.name.lower() in lowered:
            return style
    return None


def set_landscape_all_sections(doc: Document):
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        width = section.page_width
        height = section.page_height
        if width < height:
            section.page_width = height
            section.page_height = width


def parse_txt(text: str):
    lines = [line.strip() for line in text.splitlines()]

    # Убираем только пустые строки по краям, но внутри сохраняем логику
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()

    title_lines = []
    shlokas = []

    i = 0
    first_shloka_found = False

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # До первой шлоки собираем заголовок гимна
        if not first_shloka_found:
            # Пропускаем отдельные цифровые строки вроде १
            if is_number_line(line):
                i += 1
                continue

            # Шлока начинается, когда после номера/пустот идут 2 строки деванагари
            if is_devanagari(line):
                j = i + 1
                while j < len(lines) and (not lines[j].strip() or is_number_line(lines[j].strip())):
                    j += 1

                if j < len(lines) and is_devanagari(lines[j].strip()):
                    first_shloka_found = True
                else:
                    title_lines.append(line)
                    i += 1
                    continue
            else:
                i += 1
                continue

        # Пропускаем отдельную строку с номером шлоки
        if is_number_line(line):
            i += 1
            continue

        # Ищем первую строку деванагари шлоки
        if not is_devanagari(line):
            i += 1
            continue

        dev1 = line

        # Вторая строка деванагари
        i += 1
        while i < len(lines) and (not lines[i].strip() or is_number_line(lines[i].strip())):
            i += 1

        if i >= len(lines) or not is_devanagari(lines[i].strip()):
            # Это не шлока, пропускаем
            continue

        dev2 = lines[i].strip()

        # Ищем 2 строки IAST
        i += 1
        while i < len(lines) and (not lines[i].strip() or not is_iast(lines[i].strip())):
            i += 1

        if i >= len(lines):
            break
        iast1 = lines[i].strip()

        i += 1
        while i < len(lines) and not lines[i].strip():
            i += 1

        if i >= len(lines) or not is_iast(lines[i].strip()):
            break
        iast2 = lines[i].strip()

        # Перевод: один абзац до следующего пустого блока/разбора
        i += 1
        while i < len(lines) and not lines[i].strip():
            i += 1

        if i >= len(lines):
            trans = ""
            gloss = ""
            shlokas.append((dev1, dev2, iast1, iast2, trans, gloss))
            break

        trans_lines = []
        while i < len(lines):
            cur = lines[i].strip()
            if not cur:
                if trans_lines:
                    break
                i += 1
                continue
            # Если сразу начинается новая деванагари-строка — перевод пустой
            if is_devanagari(cur):
                break
            # Если строка похожа на пословный разбор: много скобок
            if "(" in cur and ")" in cur and trans_lines:
                break
            trans_lines.append(cur)
            i += 1

        trans = " ".join(trans_lines).strip()

        # Пропускаем пустые строки перед разбором
        while i < len(lines) and not lines[i].strip():
            i += 1

        gloss_lines = []
        while i < len(lines):
            cur = lines[i].strip()
            if not cur:
                if gloss_lines:
                    # Смотрим, не начинается ли после пустоты новая шлока
                    j = i + 1
                    while j < len(lines) and (not lines[j].strip() or is_number_line(lines[j].strip())):
                        j += 1
                    if j < len(lines) and is_devanagari(lines[j].strip()):
                        break
                i += 1
                continue

            if is_number_line(cur):
                j = i + 1
                while j < len(lines) and (not lines[j].strip() or is_number_line(lines[j].strip())):
                    j += 1
                if j < len(lines) and is_devanagari(lines[j].strip()):
                    break

            if is_devanagari(cur):
                break

            gloss_lines.append(cur)
            i += 1

        gloss = " ".join(gloss_lines).strip()

        shlokas.append((dev1, dev2, iast1, iast2, trans, gloss))

    return title_lines, shlokas


def add_runs_with_italic_parentheses(paragraph, text: str):
    parts = re.split(r"(\([^)]*\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("(") and part.endswith(")"):
            paragraph.add_run("(")
            r = paragraph.add_run(part[1:-1])
            r.italic = True
            paragraph.add_run(")")
        else:
            paragraph.add_run(part)


def build_docx(template_path: str, title_lines, shlokas, output_path: str):
    doc = Document(template_path)

    # Ставим альбомную ориентацию у всех секций ДО и ПОСЛЕ пересборки
    set_landscape_all_sections(doc)

    # Границы: шапку до первого devanagari не трогаем;
    # footer от "Занятия ведет" и далее тоже не трогаем
    header_end = 0
    for idx, p in enumerate(doc.paragraphs):
        if p.style and p.style.name and p.style.name.lower() == "devanagari":
            header_end = idx
            break

    footer_start = len(doc.paragraphs)
    for idx, p in enumerate(doc.paragraphs):
        if "Занятия ведет" in p.text:
            footer_start = idx
            break

    header_xml = [copy.deepcopy(p._element) for p in doc.paragraphs[:header_end]]
    footer_xml = [copy.deepcopy(p._element) for p in doc.paragraphs[footer_start:]]

    styles = {
        "devanagari": get_style(doc, ["devanagari"]),
        "iast": get_style(doc, ["iast"]),
        "trans": get_style(doc, ["trans-12"]),
        "gloss": get_style(doc, ["iast-14"]),
        "normal": get_style(doc, ["Normal", "Обычный"]),
    }

    if styles["normal"] is None:
        styles["normal"] = doc.styles["Normal"]

    missing = [k for k, v in styles.items() if v is None]
    if missing:
        raise RuntimeError(f"Не найдены стили в шаблоне: {', '.join(missing)}")

    body = doc._element.body
    for el in list(body):
        body.remove(el)

    for el in header_xml:
        body.append(el)

    # Ещё раз после пересборки тела
    set_landscape_all_sections(doc)

    # Добавляем заголовок гимна перед первой шлокой:
    # тем же стилем devanagari, через line break внутри одного абзаца
    if title_lines:
        p = doc.add_paragraph(style=styles["devanagari"])
        for n, line in enumerate(title_lines):
            if n > 0:
                p.add_run().add_break()
            p.add_run(line)

    total = len(shlokas)

    for idx, (dev1, dev2, iast1, iast2, trans, gloss) in enumerate(shlokas):
        p_dev = doc.add_paragraph(style=styles["devanagari"])
        p_dev.add_run(dev1)
        p_dev.add_run().add_break()
        p_dev.add_run(dev2)

        p_iast = doc.add_paragraph(style=styles["iast"])
        p_iast.add_run(iast1)
        p_iast.add_run().add_break()
        p_iast.add_run(iast2)

        p_trans = doc.add_paragraph(style=styles["trans"])
        p_trans.add_run(trans)

        p_gloss = doc.add_paragraph(style=styles["gloss"])
        add_runs_with_italic_parentheses(p_gloss, gloss)

        # Пустая строка + разрыв страницы только МЕЖДУ шлоками, не перед footer
        if idx < total - 1:
            p_empty = doc.add_paragraph(style=styles["normal"])
            r = p_empty.add_run("")
            r.font.size = Pt(12)
            r.add_break(WD_BREAK.PAGE)

    for el in footer_xml:
        body.append(el)

    set_landscape_all_sections(doc)
    doc.save(output_path)


def main():
    logs = []

    if not os.path.exists(TEMPLATE):
        raise FileNotFoundError(f"Не найден шаблон: {TEMPLATE}")

    if not os.path.isdir(INPUT_DIR):
        raise FileNotFoundError(f"Не найдена папка input: {INPUT_DIR}")

    txt_files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".txt")]
    if not txt_files:
        raise FileNotFoundError("В папке input нет .txt файлов")

    for fname in txt_files:
        in_path = os.path.join(INPUT_DIR, fname)

        with open(in_path, "r", encoding="utf-8") as f:
            text = f.read()

        title_lines, shlokas = parse_txt(text)

        if not shlokas:
            logs.append(f"[ERROR] {fname}: не распознано ни одной шлоки")
            continue

        out_name = os.path.splitext(fname)[0] + ".docx"
        out_path = os.path.join(OUTPUT_DIR, out_name)

        build_docx(TEMPLATE, title_lines, shlokas, out_path)
        logs.append(f"[OK] {fname}: {len(shlokas)} шлок")

    with open(LOG_FILE, "w", encoding="utf-8") as f:
        f.write("\n".join(logs))

    print("Готово.")


if __name__ == "__main__":
    main()